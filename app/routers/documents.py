# app/routers/documents.py
import asyncio
import base64
from io import BytesIO
import mimetypes
from pathlib import Path
import logging
from uuid import UUID

from fastapi import APIRouter, BackgroundTasks, Body, Depends, File, HTTPException, Query, UploadFile
from fastapi.responses import FileResponse
from sqlalchemy import func, select
from sqlalchemy.ext.asyncio import AsyncSession

from app.config import (
    DOCUMENT_STORAGE_DIR,
    MAX_DOCUMENT_SIZE_BYTES,
    OCR_DPI,
    OCR_FALLBACK_ENABLED,
    OCR_FORCE_ALL_PAGES,
    OCR_LANGUAGES,
    OCR_MIN_NATIVE_QUALITY,
    OCR_MIN_TEXT_CHARS,
    OCR_PREPROCESS_BIN_THRESHOLD,
    OCR_PREPROCESS_BIN_THRESHOLDS,
    OCR_PREPROCESS_ENABLED,
    OCR_PREPROCESS_UPSCALE,
    OCR_QUALITY_MARGIN,
    OCR_TIMEOUT_SECONDS,
    OCR_TESSERACT_CALL_TIMEOUT_SECONDS,
    OCR_TESSERACT_OEM,
    OCR_TESSERACT_PSMS,
    OCR_USE_VISION_API,
    OPENAI_API_KEY,
    POPPLER_PATH,
    TESSERACT_CMD,
)
from app.db.connection import AsyncSessionLocal, get_db
from app.db.models import AssessmentRun, BusinessUnit, Document, Extraction, User

# Límite de OCRs concurrentes para no agotar la memoria en la máquina de fly.io (512 MB)
_ocr_semaphore = asyncio.Semaphore(1)
from app.dependencies.auth import AuthContext, get_bu_auth_context, require_bu_roles_with_audit
from app.schemas.assessment import AssessmentRunRead
from app.schemas.document import DocumentListResponse, DocumentRead
from app.schemas.extraction import ExtractionRead
from app.services.document_storage import DocumentStorageService
from app.services.run_enricher import enrich_assessment_runs
from app.services.usage_service import DOC_DELETED, DOC_UPLOADED, track_event

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/documents", tags=["documents"])
storage_service = DocumentStorageService(DOCUMENT_STORAGE_DIR)


def _doc_to_read(doc: Document, bu_code: str | None, created_by_name: str | None = None) -> DocumentRead:
    return DocumentRead(
        id=doc.id, bu_id=doc.bu_id, bu_code=bu_code,
        created_by_name=created_by_name,
        title=doc.title, filename=doc.filename, mime_type=doc.mime_type,
        size_bytes=doc.size_bytes, storage_key=doc.storage_key,
        created_by=doc.created_by, created_at=doc.created_at,
        status=doc.status, ocr_text=doc.ocr_text, ocr_error=doc.ocr_error,
    )


async def _enrich_docs(docs: list[Document], db: AsyncSession) -> list[DocumentRead]:
    bu_ids = {d.bu_id for d in docs}
    bu_codes: dict = {}
    if bu_ids:
        result = await db.execute(select(BusinessUnit).where(BusinessUnit.id.in_(bu_ids)))
        for bu in result.scalars().all():
            bu_codes[bu.id] = bu.code

    user_ids = {d.created_by for d in docs if d.created_by}
    user_names: dict = {}
    if user_ids:
        result = await db.execute(select(User).where(User.id.in_(user_ids)))
        for u in result.scalars().all():
            user_names[u.id] = u.full_name or u.email

    return [_doc_to_read(d, bu_codes.get(d.bu_id), user_names.get(d.created_by)) for d in docs]


def _parse_int_csv(raw: str, default_values: list[int]) -> list[int]:
    values: list[int] = []
    for chunk in (raw or "").split(","):
        chunk = chunk.strip()
        if not chunk:
            continue
        try:
            values.append(int(chunk))
        except ValueError:
            continue
    return values or default_values


def _build_ocr_images(image, image_ops, image_filter):
    """
    Produce varias variantes de imagen para OCR y mejorar recall en escaneos difíciles.
    """
    if not OCR_PREPROCESS_ENABLED:
        return [image]

    gray = image_ops.grayscale(image)
    if OCR_PREPROCESS_UPSCALE > 1.0:
        width = max(1, int(gray.width * OCR_PREPROCESS_UPSCALE))
        height = max(1, int(gray.height * OCR_PREPROCESS_UPSCALE))
        gray = gray.resize((width, height))

    denoised = gray.filter(image_filter.MedianFilter(size=3))
    sharpened = denoised.filter(image_filter.SHARPEN)
    contrasted = image_ops.autocontrast(sharpened)

    thresholds = _parse_int_csv(OCR_PREPROCESS_BIN_THRESHOLDS, [OCR_PREPROCESS_BIN_THRESHOLD])
    variants = [contrasted]
    for threshold in thresholds:
        threshold = max(1, min(254, threshold))
        variants.append(
            contrasted.point(
                lambda px, t=threshold: 255 if px > t else 0,
                mode="1",
            )
        )
    return variants


def _ocr_best_text_for_page(variants, pytesseract_module) -> str:
    """
    Ejecuta OCR con varios PSM/variantes y elige el texto de mejor calidad estimada.
    """
    psms = _parse_int_csv(OCR_TESSERACT_PSMS, [6])
    best_text = ""
    best_score = -1.0

    for variant in variants:
        for psm in psms:
            try:
                config = f"--oem {max(0, OCR_TESSERACT_OEM)} --psm {max(0, psm)} -c preserve_interword_spaces=1"
                candidate = (
                    pytesseract_module.image_to_string(
                        variant,
                        lang=OCR_LANGUAGES,
                        config=config,
                        timeout=max(1.0, OCR_TESSERACT_CALL_TIMEOUT_SECONDS),
                    )
                    or ""
                ).strip()
            except Exception:
                continue

            # Score quality and slightly reward longer coherent outputs for invoices.
            score = _text_quality_score(candidate)
            score += min(len(candidate), 2500) / 2500 * 0.08
            if score > best_score:
                best_score = score
                best_text = candidate

    return best_text


async def _run_ocr_background(document_id: UUID, content: bytes, filename: str) -> None:
    """
    Tarea de background: ejecuta OCR sobre el contenido del documento y
    actualiza el estado en BD. Se lanza tras el upload y no bloquea la respuesta.
    El semáforo limita los OCR concurrentes para no agotar la RAM (512 MB en fly.io).
    """
    async with _ocr_semaphore:
        async with AsyncSessionLocal() as db:
            try:
                document = await db.get(Document, document_id)
                if not document:
                    return

                document.status = "processing"
                db.add(document)
                await db.commit()

                suffix = Path(filename).suffix.lower()
                try:
                    if suffix in {".txt", ".md", ".json", ".csv", ".xml", ".html"}:
                        text = _decode_text_bytes(content)
                    elif suffix == ".pdf":
                        native_pages, has_images = await asyncio.gather(
                            asyncio.to_thread(_extract_pdf_text_native_pages, content),
                            asyncio.to_thread(_pdf_has_embedded_images, content),
                        )
                        native_text = "\n\n".join(p for p in native_pages if p).strip()
                        text = native_text
                        should_run_ocr = OCR_FORCE_ALL_PAGES or has_images or _should_use_ocr_on_any_page(native_pages)
                        use_vision = has_images and OCR_USE_VISION_API
                        if OCR_FALLBACK_ENABLED and should_run_ocr:
                            try:
                                if use_vision:
                                    ocr_pages = await asyncio.wait_for(
                                        _extract_pdf_pages_vision(content),
                                        timeout=max(30, OCR_TIMEOUT_SECONDS),
                                    )
                                else:
                                    ocr_pages = await asyncio.wait_for(
                                        asyncio.to_thread(_extract_pdf_text_ocr_pages, content),
                                        timeout=max(5, OCR_TIMEOUT_SECONDS),
                                    )
                                text, _ = _combine_best_pdf_pages(native_pages, ocr_pages, prefer_ocr=has_images)
                            except Exception as ocr_err:
                                logger.warning("OCR background fallo para %s: %s", filename, ocr_err)
                                text = native_text
                    elif suffix == ".docx":
                        text = await asyncio.to_thread(_extract_docx_text, content)
                    else:
                        raise ValueError(f"Formato no soportado: {suffix}")

                    document = await db.get(Document, document_id)
                    if document:
                        document.status = "processed"
                        document.ocr_text = text.strip() if text else ""
                        db.add(document)
                        await db.commit()

                except Exception as exc:
                    logger.exception("OCR background fallo para documento %s", document_id)
                    document = await db.get(Document, document_id)
                    if document:
                        document.status = "failed"
                        document.ocr_error = str(exc)[:500]
                        db.add(document)
                        await db.commit()

            except Exception:
                logger.exception("Error critico en background OCR para documento %s", document_id)


@router.post("/", response_model=DocumentRead, status_code=201)
async def upload_document(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    db: AsyncSession = Depends(get_db),
    auth: AuthContext = Depends(get_bu_auth_context),
):
    await require_bu_roles_with_audit(
        auth,
        {"admin_global", "bu_admin", "bu_user"},
        "No tienes permisos para subir documentos en esta BU",
        db,
        action="document.upload",
        resource_type="document",
    )

    filename = file.filename or "document"
    content = await file.read()
    if not content:
        raise HTTPException(status_code=400, detail="El archivo está vacío")

    if len(content) > MAX_DOCUMENT_SIZE_BYTES:
        size_mb = MAX_DOCUMENT_SIZE_BYTES / (1024 * 1024)
        raise HTTPException(
            status_code=413,
            detail=f"Archivo demasiado grande. Máximo permitido: {size_mb:.1f} MB",
        )

    bu = await db.get(BusinessUnit, auth.bu_id)
    if not bu or not bu.is_active:
        raise HTTPException(status_code=404, detail="BU no encontrada o inactiva")

    raw_mime = file.content_type or "application/octet-stream"
    if raw_mime in ("application/octet-stream", "binary/octet-stream", ""):
        guessed, _ = mimetypes.guess_type(filename)
        if guessed:
            raw_mime = guessed

    storage_key, digest = storage_service.save(content, filename)
    document = Document(
        bu_id=auth.bu_id,
        title=filename,
        filename=filename,
        mime_type=raw_mime,
        size_bytes=len(content),
        sha256=digest,
        storage_key=storage_key,
        created_by=auth.actor_user_id,
        status="pending",
    )
    db.add(document)
    await db.flush()

    # Registrar evento de uso en la misma transacción
    await track_event(
        db,
        bu_id=auth.bu_id,
        event_type=DOC_UPLOADED,
        user_id=auth.actor_user_id,
        metadata={"filename": filename, "mime_type": raw_mime, "size_bytes": len(content)},
    )

    await db.commit()
    await db.refresh(document)

    background_tasks.add_task(_run_ocr_background, document.id, content, filename)

    return _doc_to_read(document, bu.code)


@router.post("/from-base64", response_model=DocumentRead, status_code=201)
async def upload_document_base64(
    background_tasks: BackgroundTasks,
    payload: dict = Body(...),
    db: AsyncSession = Depends(get_db),
    auth: AuthContext = Depends(get_bu_auth_context),
):
    await require_bu_roles_with_audit(
        auth,
        {"admin_global", "bu_admin", "bu_user"},
        "No tienes permisos para subir documentos en esta BU",
        db,
        action="document.upload",
        resource_type="document",
    )

    filename = str(payload.get("filename") or "document")
    b64 = payload.get("content_base64") or ""
    if not b64:
        raise HTTPException(status_code=400, detail="content_base64 es obligatorio")
    try:
        content = base64.b64decode(b64)
    except Exception:
        raise HTTPException(status_code=400, detail="content_base64 no es base64 válido")

    if not content:
        raise HTTPException(status_code=400, detail="El archivo está vacío")
    if len(content) > MAX_DOCUMENT_SIZE_BYTES:
        size_mb = MAX_DOCUMENT_SIZE_BYTES / (1024 * 1024)
        raise HTTPException(status_code=413, detail=f"Archivo demasiado grande. Máximo: {size_mb:.1f} MB")

    bu = await db.get(BusinessUnit, auth.bu_id)
    if not bu or not bu.is_active:
        raise HTTPException(status_code=404, detail="BU no encontrada o inactiva")

    guessed, _ = mimetypes.guess_type(filename)
    raw_mime = guessed or "application/octet-stream"

    storage_key, digest = storage_service.save(content, filename)
    document = Document(
        bu_id=auth.bu_id,
        title=filename,
        filename=filename,
        mime_type=raw_mime,
        size_bytes=len(content),
        sha256=digest,
        storage_key=storage_key,
        created_by=auth.actor_user_id,
        status="pending",
    )
    db.add(document)
    await db.flush()

    # Registrar evento de uso en la misma transacción
    await track_event(
        db,
        bu_id=auth.bu_id,
        event_type=DOC_UPLOADED,
        user_id=auth.actor_user_id,
        metadata={"filename": filename, "mime_type": raw_mime, "size_bytes": len(content)},
    )

    await db.commit()
    await db.refresh(document)

    background_tasks.add_task(_run_ocr_background, document.id, content, filename)

    return _doc_to_read(document, bu.code)


@router.get("/", response_model=DocumentListResponse)
async def list_documents(
    limit: int = 20,
    offset: int = 0,
    db: AsyncSession = Depends(get_db),
    auth: AuthContext = Depends(get_bu_auth_context),
):
    limit = max(1, min(limit, 200))
    offset = max(0, offset)

    count_stmt = select(func.count(Document.id)).where(Document.bu_id == auth.bu_id)
    total = (await db.execute(count_stmt)).scalar_one()

    stmt = (
        select(Document)
        .where(Document.bu_id == auth.bu_id)
        .order_by(Document.created_at.desc())
        .offset(offset)
        .limit(limit)
    )
    result = await db.execute(stmt)
    items = list(result.scalars().all())
    enriched = await _enrich_docs(items, db)
    return DocumentListResponse(items=enriched, total=total)


@router.get("/{document_id}", response_model=DocumentRead)
async def get_document(
    document_id: UUID,
    db: AsyncSession = Depends(get_db),
    auth: AuthContext = Depends(get_bu_auth_context),
):
    result = await db.execute(
        select(Document).where(Document.id == document_id, Document.bu_id == auth.bu_id)
    )
    document = result.scalars().first()
    if not document:
        raise HTTPException(status_code=404, detail="Documento no encontrado")
    enriched = await _enrich_docs([document], db)
    return enriched[0]


@router.get("/{document_id}/download")
async def download_document(
    document_id: UUID,
    db: AsyncSession = Depends(get_db),
    auth: AuthContext = Depends(get_bu_auth_context),
):
    result = await db.execute(
        select(Document).where(Document.id == document_id, Document.bu_id == auth.bu_id)
    )
    document = result.scalars().first()
    if not document:
        raise HTTPException(status_code=404, detail="Documento no encontrado")

    file_path = storage_service.resolve(document.storage_key)
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="Archivo no encontrado en storage")

    return FileResponse(path=file_path, filename=document.filename, media_type=document.mime_type)


@router.delete("/{document_id}", status_code=204)
async def delete_document(
    document_id: UUID,
    db: AsyncSession = Depends(get_db),
    auth: AuthContext = Depends(get_bu_auth_context),
):
    await require_bu_roles_with_audit(
        auth,
        {"admin_global", "bu_admin", "bu_user"},
        "No tienes permisos para eliminar documentos en esta BU",
        db,
        action="document.delete",
        resource_type="document",
        resource_id=str(document_id),
    )

    result = await db.execute(
        select(Document).where(Document.id == document_id, Document.bu_id == auth.bu_id)
    )
    document = result.scalars().first()
    if not document:
        raise HTTPException(status_code=404, detail="Documento no encontrado")

    # Registrar evento antes del delete (necesitamos los datos del documento)
    await track_event(
        db,
        bu_id=auth.bu_id,
        event_type=DOC_DELETED,
        user_id=auth.actor_user_id,
        metadata={
            "document_id": str(document_id),
            "filename": document.filename,
            "size_bytes": document.size_bytes,
        },
    )

    storage_service.delete(document.storage_key)
    await db.delete(document)
    await db.commit()


@router.get("/{document_id}/runs", response_model=list[ExtractionRead])
async def list_document_runs(
    document_id: UUID,
    limit: int = 50,
    db: AsyncSession = Depends(get_db),
    auth: AuthContext = Depends(get_bu_auth_context),
):
    limit = max(1, min(limit, 200))

    document_result = await db.execute(
        select(Document).where(Document.id == document_id, Document.bu_id == auth.bu_id)
    )
    document = document_result.scalars().first()
    if not document:
        raise HTTPException(status_code=404, detail="Documento no encontrado")

    runs_result = await db.execute(
        select(Extraction)
        .where(
            Extraction.document_id == document_id,
            Extraction.bu_id == auth.bu_id,
        )
        .order_by(Extraction.created_at.desc())
        .limit(limit)
    )
    return list(runs_result.scalars().all())


@router.get("/{document_id}/assessment-runs", response_model=list[AssessmentRunRead])
async def list_document_assessment_runs(
    document_id: UUID,
    limit: int = Query(20, ge=1, le=100),
    offset: int = Query(0, ge=0),
    db: AsyncSession = Depends(get_db),
    auth: AuthContext = Depends(get_bu_auth_context),
):
    doc_result = await db.execute(
        select(Document).where(Document.id == document_id, Document.bu_id == auth.bu_id)
    )
    if not doc_result.scalars().first():
        raise HTTPException(status_code=404, detail="Documento no encontrado")
    runs_result = await db.execute(
        select(AssessmentRun)
        .where(AssessmentRun.document_id == document_id, AssessmentRun.bu_id == auth.bu_id)
        .order_by(AssessmentRun.created_at.desc())
        .offset(offset)
        .limit(limit)
    )
    runs = list(runs_result.scalars().all())
    return await enrich_assessment_runs(runs, db)


def _decode_text_bytes(content: bytes) -> str:
    """
    Decodifica bytes de texto intentando UTF-8 y fallback a latin-1.
    """
    try:
        return content.decode("utf-8")
    except UnicodeDecodeError:
        return content.decode("latin-1")


def _extract_pdf_text_native_pages(content: bytes) -> list[str]:
    """
    Extrae texto nativo por página para poder decidir OCR en documentos mixtos.
    """
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise HTTPException(
            status_code=400,
            detail="Para procesar PDF instala pypdf: pip install pypdf",
        ) from exc

    reader = PdfReader(BytesIO(content))
    return [(page.extract_text() or "").strip() for page in reader.pages]


def _pdf_has_embedded_images(content: bytes) -> bool:
    """
    Detecta si el PDF contiene imágenes embebidas en alguna página (PDF mixto).
    En PDFs mixtos el texto nativo es incompleto: datos clave pueden estar
    sólo en la imagen, por lo que es necesario forzar OCR.
    """
    try:
        from pypdf import PdfReader
        reader = PdfReader(BytesIO(content))
        for page in reader.pages:
            if page.images:
                return True
    except Exception:
        pass
    return False


def _encode_page_to_b64(image) -> str:
    import base64
    buf = BytesIO()
    image.save(buf, format="JPEG", quality=85)
    return base64.b64encode(buf.getvalue()).decode()


async def _extract_pdf_pages_vision(content: bytes) -> list[str]:
    """
    Extrae texto de cada página del PDF usando GPT-4o Vision.
    Páginas procesadas en paralelo (semáforo de 4) para maximizar velocidad.
    Conversión PDF→imagen en thread para no bloquear el event loop.
    """
    try:
        from pdf2image import convert_from_bytes
    except ImportError as exc:
        raise HTTPException(
            status_code=400,
            detail="Para Vision OCR instala pdf2image: pip install pdf2image",
        ) from exc

    import httpx

    try:
        poppler_path = POPPLER_PATH or None
        images = await asyncio.to_thread(
            convert_from_bytes, content, dpi=150, poppler_path=poppler_path
        )
    except Exception as exc:
        raise HTTPException(
            status_code=400,
            detail="No se pudo convertir PDF a imagen para Vision OCR.",
        ) from exc

    api_headers = {
        "Authorization": f"Bearer {OPENAI_API_KEY}",
        "Content-Type": "application/json",
    }
    sem = asyncio.Semaphore(4)

    async def _process_page(client: "httpx.AsyncClient", image, page_idx: int) -> str:
        async with sem:
            try:
                b64 = await asyncio.to_thread(_encode_page_to_b64, image)
                payload = {
                    "model": "gpt-4o",
                    "max_tokens": 2048,
                    "messages": [{
                        "role": "user",
                        "content": [
                            {
                                "type": "text",
                                "text": (
                                    "Extrae todo el texto visible en esta página de documento. "
                                    "Incluye todos los campos: nombres, fechas, importes, NIFs/CIFs, "
                                    "direcciones y cualquier texto en márgenes o pies de página. "
                                    "Devuelve SOLO el texto extraído, sin explicaciones."
                                ),
                            },
                            {
                                "type": "image_url",
                                "image_url": {"url": f"data:image/jpeg;base64,{b64}", "detail": "high"},
                            },
                        ],
                    }],
                }
                resp = await client.post(
                    "https://api.openai.com/v1/chat/completions",
                    headers=api_headers,
                    json=payload,
                )
                resp.raise_for_status()
                return (resp.json()["choices"][0]["message"]["content"] or "").strip()
            except Exception as exc:
                logger.warning("Vision API fallo en página %d: %s", page_idx, exc)
                return ""

    async with httpx.AsyncClient(timeout=90.0) as client:
        results = await asyncio.gather(
            *[_process_page(client, img, i) for i, img in enumerate(images)]
        )

    return list(results)


def _extract_pdf_text_ocr_pages(content: bytes) -> list[str]:
    """
    Extrae texto OCR por página para comparar contra extracción nativa página a página.
    """
    try:
        from pdf2image import convert_from_bytes
        from PIL import ImageFilter, ImageOps
        import pytesseract
    except ImportError as exc:
        raise HTTPException(
            status_code=400,
            detail=(
                "OCR no disponible. Instala dependencias Python: "
                "pip install pdf2image pytesseract pillow"
            ),
        ) from exc

    if TESSERACT_CMD:
        pytesseract.pytesseract.tesseract_cmd = TESSERACT_CMD

    try:
        poppler_path = POPPLER_PATH or None
        images = convert_from_bytes(content, dpi=OCR_DPI, poppler_path=poppler_path)
    except Exception as exc:
        raise HTTPException(
            status_code=400,
            detail=(
                "No se pudo convertir PDF a imagen para OCR. "
                "En Windows instala Poppler y agrega su bin al PATH."
            ),
        ) from exc

    chunks = []
    for image in images:
        try:
            variants = _build_ocr_images(image, ImageOps, ImageFilter)
            text = _ocr_best_text_for_page(variants, pytesseract)
        except Exception as exc:
            raise HTTPException(
                status_code=400,
                detail=(
                    "Error ejecutando OCR con Tesseract. "
                    "Verifica que Tesseract este instalado y disponible en PATH."
                ),
            ) from exc
        chunks.append((text or "").strip())

    return chunks


def _should_use_ocr(native_text: str) -> bool:
    """
    Decide si activar OCR cuando el texto embebido es insuficiente.
    """
    cleaned = (native_text or "").strip()
    if len(cleaned) < OCR_MIN_TEXT_CHARS:
        return True
    return _text_quality_score(cleaned) < OCR_MIN_NATIVE_QUALITY


def _text_quality_score(text: str) -> float:
    """
    Estima calidad de texto extraído (0-1): penaliza texto corrupto y caracteres de reemplazo.
    """
    cleaned = (text or "").strip()
    if not cleaned:
        return 0.0

    total = len(cleaned)
    printable_ratio = sum(1 for ch in cleaned if ch.isprintable()) / total
    # Readable: permite s\u00edmbolos v\u00e1lidos (\u20ac, %, /, ., ,) \u2014 s\u00f3lo excluye caracteres de control
    readable_ratio = sum(1 for ch in cleaned if ch >= " " or ch in "\n\r\t") / total
    replacement_ratio = cleaned.count("\ufffd") / total

    score = (0.6 * printable_ratio) + (0.4 * readable_ratio) - (1.5 * replacement_ratio)
    return max(0.0, min(1.0, score))


def _pick_best_pdf_text(native_text: str, ocr_text: str, prefer_ocr: bool = False) -> str:
    """
    Elige entre texto nativo y OCR priorizando calidad legible.
    Si prefer_ocr=True (PDF mixto con imágenes), se usa OCR siempre que su
    calidad sea aceptable porque captura tanto el texto nativo como el contenido
    de las imágenes embebidas.
    """
    native = (native_text or "").strip()
    ocr = (ocr_text or "").strip()

    if not ocr:
        return native
    if not native:
        return ocr

    native_score = _text_quality_score(native)
    ocr_score = _text_quality_score(ocr)

    if prefer_ocr:
        # PDF mixto: OCR ve todo (texto + imágenes). Usarlo salvo que su calidad
        # sea claramente peor que el texto nativo.
        if ocr_score >= native_score - 0.1:
            return ocr
        return native

    if ocr_score >= native_score + OCR_QUALITY_MARGIN:
        return ocr

    if ocr_score > native_score and len(ocr) >= int(len(native) * 0.85):
        return ocr

    return native


def _should_use_ocr_on_any_page(native_pages: list[str]) -> bool:
    """
    Activa OCR si al menos una página parece pobre o con extracción deficiente.
    """
    if not native_pages:
        return True

    for page_text in native_pages:
        if _should_use_ocr(page_text):
            return True

    return False


def _combine_best_pdf_pages(
    native_pages: list[str],
    ocr_pages: list[str],
    prefer_ocr: bool = False,
) -> tuple[str, int]:
    """
    Devuelve texto final combinando mejor opción por página y cuántas páginas usaron OCR.
    prefer_ocr=True (PDF mixto): concatena nativo + OCR para maximizar recall.
    El texto nativo aporta datos estructurados con precisión; el OCR aporta contenido
    de imágenes y elementos decorativos (ej. NIF del emisor en pie de página).
    """
    max_len = max(len(native_pages), len(ocr_pages))
    merged_pages: list[str] = []
    pages_using_ocr = 0

    for idx in range(max_len):
        native_page = (native_pages[idx] if idx < len(native_pages) else "").strip()
        ocr_page = (ocr_pages[idx] if idx < len(ocr_pages) else "").strip()

        if OCR_FORCE_ALL_PAGES and ocr_page:
            chosen = ocr_page
            pages_using_ocr += 1
        elif prefer_ocr and native_page and ocr_page:
            # PDF mixto: combinar ambas fuentes — el LLM puede trabajar con
            # información redundante y extraer de la fuente que la contenga.
            chosen = native_page + "\n\n" + ocr_page
            pages_using_ocr += 1
        else:
            chosen = _pick_best_pdf_text(native_page, ocr_page, prefer_ocr=prefer_ocr).strip()
            if chosen == ocr_page and chosen:
                pages_using_ocr += 1

        merged_pages.append(chosen)

    return "\n\n".join(page for page in merged_pages if page).strip(), pages_using_ocr


def _extract_docx_text(content: bytes) -> str:
    """
    Extrae texto de DOCX usando python-docx si está disponible.
    """
    try:
        from docx import Document
    except ImportError as exc:
        raise HTTPException(
            status_code=400,
            detail="Para procesar DOCX instala python-docx: pip install python-docx",
        ) from exc

    document = Document(BytesIO(content))
    parts = [p.text for p in document.paragraphs if p.text]
    for table in document.tables:
        for row in table.rows:
            row_text = "  ".join(cell.text for cell in row.cells if cell.text.strip())
            if row_text.strip():
                parts.append(row_text)
    return "\n".join(parts).strip()


@router.post("/parse")
async def parse_document(
    file: UploadFile = File(...),
    db: AsyncSession = Depends(get_db),
    auth: AuthContext = Depends(get_bu_auth_context),
):
    """
    Recibe un archivo y devuelve texto plano para reutilizarlo en /extract.
    Soporta txt/md/json/csv/xml/html y, opcionalmente, pdf/docx.
    """
    await require_bu_roles_with_audit(
        auth,
        {"admin_global", "bu_admin", "bu_user"},
        "No tienes permisos para procesar documentos en esta BU",
        db,
        action="document.parse",
        resource_type="document",
    )

    filename = file.filename or "document"
    suffix = Path(filename).suffix.lower()

    content = await file.read()
    if not content:
        raise HTTPException(status_code=400, detail="El archivo está vacío")
    
    # Validar tamaño del archivo
    if len(content) > MAX_DOCUMENT_SIZE_BYTES:
        size_mb = MAX_DOCUMENT_SIZE_BYTES / (1024 * 1024)
        raise HTTPException(
            status_code=413,
            detail=f"Archivo demasiado grande. Máximo permitido: {size_mb:.1f} MB"
        )

    used_ocr = False
    ocr_warning = None

    try:
        if suffix in {".txt", ".md", ".json", ".csv", ".xml", ".html"}:
            text = _decode_text_bytes(content)
        elif suffix == ".pdf":
            native_pages, has_images = await asyncio.gather(
                asyncio.to_thread(_extract_pdf_text_native_pages, content),
                asyncio.to_thread(_pdf_has_embedded_images, content),
            )
            native_text = "\n\n".join(page for page in native_pages if page).strip()
            text = native_text

            should_run_ocr = OCR_FORCE_ALL_PAGES or has_images or _should_use_ocr_on_any_page(native_pages)
            use_vision = has_images and OCR_USE_VISION_API
            if OCR_FALLBACK_ENABLED and should_run_ocr:
                logger.info(
                    "Documento PDF procesado con OCR por pagina (force_all_pages=%s, has_images=%s, vision=%s, chars=%s, score=%.3f).",
                    OCR_FORCE_ALL_PAGES,
                    has_images,
                    use_vision,
                    len(native_text),
                    _text_quality_score(native_text),
                )
                try:
                    if use_vision:
                        ocr_pages = await asyncio.wait_for(
                            _extract_pdf_pages_vision(content),
                            timeout=max(30, OCR_TIMEOUT_SECONDS),
                        )
                    else:
                        ocr_pages = await asyncio.wait_for(
                            asyncio.to_thread(_extract_pdf_text_ocr_pages, content),
                            timeout=max(5, OCR_TIMEOUT_SECONDS),
                        )
                    text, pages_using_ocr = _combine_best_pdf_pages(
                        native_pages, ocr_pages, prefer_ocr=has_images
                    )
                    used_ocr = pages_using_ocr > 0

                    logger.info(
                        "Seleccion por pagina completada (native_score=%.3f, final_score=%.3f, pages_using_ocr=%s/%s)",
                        _text_quality_score(native_text),
                        _text_quality_score(text),
                        pages_using_ocr,
                        max(len(native_pages), len(ocr_pages)),
                    )
                except asyncio.TimeoutError:
                    if native_text.strip():
                        text = native_text
                        used_ocr = False
                        ocr_warning = (
                            "OCR excedio el tiempo limite y se uso texto nativo del PDF. "
                            "Ajusta OCR_TIMEOUT_SECONDS u optimiza el archivo si necesitas OCR completo."
                        )
                        logger.warning(
                            "OCR timeout para %s tras %ss; usando texto nativo.",
                            filename,
                            OCR_TIMEOUT_SECONDS,
                        )
                    else:
                        raise HTTPException(
                            status_code=504,
                            detail=(
                                "OCR excedio el tiempo limite y el PDF no contiene texto nativo util. "
                                "Reduce tamano/paginas o aumenta OCR_TIMEOUT_SECONDS."
                            ),
                        )
                except HTTPException as ocr_exc:
                    if native_text.strip():
                        text = native_text
                        used_ocr = False
                        ocr_warning = (
                            "OCR no disponible en este entorno; se uso texto nativo del PDF. "
                            "Para OCR instala Tesseract y Poppler."
                        )
                        logger.warning(
                            "OCR fallido para %s; usando texto nativo. detalle=%s",
                            filename,
                            ocr_exc.detail,
                        )
                    else:
                        raise
        elif suffix == ".docx":
            text = _extract_docx_text(content)
        else:
            raise HTTPException(
                status_code=400,
                detail=(
                    "Formato no soportado. Usa txt, md, json, csv, xml, html, pdf o docx"
                ),
            )

        if not text.strip():
            raise HTTPException(
                status_code=400,
                detail="No se pudo extraer texto útil del archivo",
            )

        return {
            "filename": filename,
            "content_type": file.content_type,
            "char_count": len(text),
            "used_ocr": used_ocr,
            "ocr_warning": ocr_warning,
            "text": text,
            "preview": text[:1000],
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.exception(f"Error procesando archivo {filename}")
        raise HTTPException(
            status_code=400,
            detail=f"Error procesando archivo: {str(e)}"
        ) from e
