"""
Genera documentos Word (.docx) a partir de los ficheros Markdown de Centinell.
Uso:  venv\\Scripts\\python docs\\build_docs.py
"""
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


# ── Paleta de colores ─────────────────────────────────────────────────────────
NAVY       = RGBColor(0x1F, 0x38, 0x64)   # títulos principales
BLUE       = RGBColor(0x27, 0x5D, 0xA8)   # h2
TEAL       = RGBColor(0x2E, 0x75, 0xB6)   # h3
GRAY_DARK  = RGBColor(0x40, 0x40, 0x40)   # cuerpo
GRAY_MED   = RGBColor(0x70, 0x70, 0x70)   # subtítulo portada
TABLE_HEAD = RGBColor(0x1F, 0x38, 0x64)   # fondo cabecera tabla
CODE_BG    = RGBColor(0xF2, 0xF2, 0xF2)   # fondo bloque código
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
ACCENT     = RGBColor(0xED, 0x75, 0x25)   # naranja para acentos


# ── Helpers XML ───────────────────────────────────────────────────────────────

def _set_cell_bg(cell, color):
    """color: RGBColor instance (which is an int subclass) or hex string like 'FF0000'"""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    hex_color = f"{int(color):06X}" if isinstance(color, int) else str(color)
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _set_cell_border(cell, side="bottom", size=4, color="1F3864"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    border = OxmlElement(f"w:{side}")
    border.set(qn("w:val"),   "single")
    border.set(qn("w:sz"),    str(size))
    border.set(qn("w:space"), "0")
    border.set(qn("w:color"), color)
    tcBorders.append(border)


def _add_page_number_footer(doc):
    section = doc.sections[0]
    footer  = section.footer
    para    = footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.clear()
    run = para.add_run()
    run.font.size  = Pt(9)
    run.font.color.rgb = GRAY_MED
    fldChar = OxmlElement("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "begin")
    run._r.append(fldChar)
    instrText = OxmlElement("w:instrText")
    instrText.text = " PAGE "
    run._r.append(instrText)
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar2)


def _add_horizontal_rule(doc):
    p    = doc.add_paragraph()
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2E75B6")
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.space_before = Pt(6)


# ── Configuración global del documento ───────────────────────────────────────

def _setup_doc(doc):
    # Márgenes
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    # Estilo Normal
    normal = doc.styles["Normal"]
    normal.font.name  = "Calibri"
    normal.font.size  = Pt(10.5)
    normal.font.color.rgb = GRAY_DARK
    normal.paragraph_format.space_after  = Pt(6)
    normal.paragraph_format.space_before = Pt(0)

    # Headings
    _h(doc, "Heading 1", NAVY,  Pt(20), bold=True,  space_before=Pt(18), space_after=Pt(8))
    _h(doc, "Heading 2", BLUE,  Pt(15), bold=True,  space_before=Pt(14), space_after=Pt(6))
    _h(doc, "Heading 3", TEAL,  Pt(12), bold=True,  space_before=Pt(10), space_after=Pt(4))
    _h(doc, "Heading 4", GRAY_DARK, Pt(11), bold=True, space_before=Pt(8), space_after=Pt(4))

    _add_page_number_footer(doc)


def _h(doc, style_name, color, size, bold=True, space_before=Pt(12), space_after=Pt(4)):
    st = doc.styles[style_name]
    st.font.name       = "Calibri"
    st.font.size       = size
    st.font.bold       = bold
    st.font.color.rgb  = color
    st.paragraph_format.space_before = space_before
    st.paragraph_format.space_after  = space_after
    st.paragraph_format.keep_with_next = True


# ── Portada ───────────────────────────────────────────────────────────────────

def _add_cover(doc, title: str, subtitle: str, date: str):
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("CENTINELL")
    run.font.name  = "Calibri"
    run.font.size  = Pt(11)
    run.font.bold  = True
    run.font.color.rgb = ACCENT
    run.font.all_caps  = True
    p.paragraph_format.space_after = Pt(6)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(title)
    run2.font.name  = "Calibri"
    run2.font.size  = Pt(28)
    run2.font.bold  = True
    run2.font.color.rgb = NAVY
    p2.paragraph_format.space_after = Pt(12)

    _add_horizontal_rule(doc)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run(subtitle)
    run3.font.name  = "Calibri"
    run3.font.size  = Pt(13)
    run3.font.color.rgb = GRAY_MED
    p3.paragraph_format.space_after = Pt(24)

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p4.add_run(date)
    run4.font.name  = "Calibri"
    run4.font.size  = Pt(10)
    run4.font.color.rgb = GRAY_MED

    doc.add_page_break()


# ── Renderizado de texto inline ──────────────────────────────────────────────

def _render_inline(para, text: str, base_size=None, base_color=None):
    """Procesa **negrita**, `código` y texto normal dentro de un párrafo."""
    # Patrón: **bold** o `code`
    pattern = r"(\*\*[^*]+\*\*|`[^`]+`)"
    parts   = re.split(pattern, text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
            if base_size:  run.font.size = base_size
            if base_color: run.font.color.rgb = base_color
        elif part.startswith("`") and part.endswith("`"):
            run = para.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        else:
            if part:
                run = para.add_run(part)
                if base_size:  run.font.size = base_size
                if base_color: run.font.color.rgb = base_color


# ── Tabla Markdown ────────────────────────────────────────────────────────────

def _parse_md_table(lines: list[str]):
    rows = []
    for line in lines:
        if re.match(r"^\|[-| :]+\|$", line.strip()):
            continue
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)
    return rows


def _add_md_table(doc, rows: list[list[str]]):
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Table Grid"

    for ri, row in enumerate(rows):
        for ci, cell_text in enumerate(row):
            if ci >= ncols:
                break
            cell = table.cell(ri, ci)
            cell.text = ""
            para = cell.paragraphs[0]
            para.paragraph_format.space_before = Pt(3)
            para.paragraph_format.space_after  = Pt(3)

            # Limpiar emojis simples (⚠️ etc.)
            clean = re.sub(r"[⚠️✅❌→←↑↓┌┐└┘│├┤┬┴┼─]", "", cell_text).strip()

            if ri == 0:
                # Cabecera
                _set_cell_bg(cell, TABLE_HEAD)
                run = para.add_run(re.sub(r"\*\*([^*]+)\*\*", r"\1", clean))
                run.font.bold        = True
                run.font.color.rgb   = WHITE
                run.font.name        = "Calibri"
                run.font.size        = Pt(10)
                para.alignment       = WD_ALIGN_PARAGRAPH.CENTER
            else:
                _render_inline(para, clean)
                for r in para.runs:
                    r.font.name = "Calibri"
                    r.font.size = Pt(10)
                # Filas alternas
                if ri % 2 == 0:
                    _set_cell_bg(cell, RGBColor(0xEE, 0xF3, 0xFA))

    # Ajustar ancho de columnas
    try:
        available = Cm(15)
        col_width  = available // ncols
        for col in table.columns:
            for cell in col.cells:
                cell.width = col_width
    except Exception:
        pass

    doc.add_paragraph()


# ── Bloque de código ──────────────────────────────────────────────────────────

def _add_code_block(doc, lines: list[str]):
    text = "\n".join(lines)
    para = doc.add_paragraph()
    para.paragraph_format.left_indent   = Cm(1)
    para.paragraph_format.space_before  = Pt(4)
    para.paragraph_format.space_after   = Pt(4)
    # Fondo gris
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "F2F2F2")
    pPr.append(shd)
    run = para.add_run(text)
    run.font.name  = "Courier New"
    run.font.size  = Pt(9)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)


# ── Parser principal ──────────────────────────────────────────────────────────

def _parse_and_render(doc, md_text: str):
    lines     = md_text.split("\n")
    i         = 0
    in_list   = False

    while i < len(lines):
        line = lines[i]

        # ── Portada / primera línea (ya gestionada fuera) ──────────────────
        if line.startswith("# ") and i == 0:
            i += 1
            continue

        # ── Línea vacía ────────────────────────────────────────────────────
        if line.strip() == "":
            in_list = False
            i += 1
            continue

        # ── Separador --- ─────────────────────────────────────────────────
        if re.match(r"^-{3,}$", line.strip()):
            _add_horizontal_rule(doc)
            i += 1
            continue

        # ── Salto de página explícito (<!-- pagebreak -->) ────────────────
        if "<!-- pagebreak -->" in line.lower():
            doc.add_page_break()
            i += 1
            continue

        # ── Bloque de código ──────────────────────────────────────────────
        if line.strip().startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            _add_code_block(doc, code_lines)
            i += 1
            continue

        # ── Tabla ─────────────────────────────────────────────────────────
        if line.strip().startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            rows = _parse_md_table(table_lines)
            _add_md_table(doc, rows)
            continue

        # ── Headings ──────────────────────────────────────────────────────
        m = re.match(r"^(#{1,4})\s+(.*)", line)
        if m:
            level = len(m.group(1))
            text  = m.group(2).strip()
            # Limpiar anclajes markdown [texto](#anchor)
            text  = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
            style = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3", 4: "Heading 4"}.get(level, "Heading 3")
            doc.add_heading(text, level=level)
            in_list = False
            i += 1
            continue

        # ── Lista con viñeta ──────────────────────────────────────────────
        m_li = re.match(r"^(\s*)[-*]\s+(.*)", line)
        if m_li:
            indent = len(m_li.group(1))
            text   = m_li.group(2)
            style  = "List Bullet 2" if indent >= 2 else "List Bullet"
            try:
                para = doc.add_paragraph(style=style)
            except Exception:
                para = doc.add_paragraph()
                para.paragraph_format.left_indent = Cm(1)
            _render_inline(para, text)
            for run in para.runs:
                run.font.name = "Calibri"
                run.font.size = Pt(10.5)
            in_list = True
            i += 1
            continue

        # ── Lista numerada ────────────────────────────────────────────────
        m_num = re.match(r"^\d+\.\s+(.*)", line)
        if m_num:
            text = m_num.group(1)
            try:
                para = doc.add_paragraph(style="List Number")
            except Exception:
                para = doc.add_paragraph()
                para.paragraph_format.left_indent = Cm(1)
            _render_inline(para, text)
            for run in para.runs:
                run.font.name = "Calibri"
                run.font.size = Pt(10.5)
            in_list = True
            i += 1
            continue

        # ── Blockquote > ──────────────────────────────────────────────────
        m_bq = re.match(r"^>\s*(.*)", line)
        if m_bq:
            text = m_bq.group(1)
            para = doc.add_paragraph()
            para.paragraph_format.left_indent  = Cm(1.2)
            para.paragraph_format.space_before = Pt(4)
            para.paragraph_format.space_after  = Pt(4)
            pPr = para._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            left = OxmlElement("w:left")
            left.set(qn("w:val"),   "single")
            left.set(qn("w:sz"),    "12")
            left.set(qn("w:space"), "4")
            left.set(qn("w:color"), "2E75B6")
            pBdr.append(left)
            pPr.append(pBdr)
            _render_inline(para, text, base_color=BLUE)
            for run in para.runs:
                run.font.name   = "Calibri"
                run.font.size   = Pt(10.5)
                run.font.italic = True
            i += 1
            continue

        # ── Párrafo normal ────────────────────────────────────────────────
        para = doc.add_paragraph()
        _render_inline(para, line)
        for run in para.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(10.5)
        in_list = False
        i += 1


# ── Función principal ─────────────────────────────────────────────────────────

def build_doc(md_path: Path, out_path: Path, title: str, subtitle: str, date: str):
    md_text = md_path.read_text(encoding="utf-8")

    # Quitar la primera línea del md (el # título) — ya va en la portada
    lines = md_text.split("\n")
    if lines and lines[0].startswith("# "):
        lines = lines[1:]
    md_body = "\n".join(lines)

    doc = Document()
    _setup_doc(doc)
    _add_cover(doc, title, subtitle, date)
    _parse_and_render(doc, md_body)

    doc.save(str(out_path))
    print(f"OK  {out_path.name}  ({out_path.stat().st_size // 1024} KB)")


# ── Entry point ───────────────────────────────────────────────────────────────

if __name__ == "__main__":
    base = Path(__file__).parent

    build_doc(
        md_path  = base / "FUNCTIONAL_SPEC.md",
        out_path = base / "Centinell_Especificacion_Funcional.docx",
        title    = "Especificación Funcional",
        subtitle = "Documentación técnica y referencia completa de API",
        date     = "Mayo 2026  ·  v1.0",
    )

    build_doc(
        md_path  = base / "COMMERCIAL_PLAN.md",
        out_path = base / "Centinell_Plan_Comercial.docx",
        title    = "Plan Comercial",
        subtitle = "Tarifas, márgenes, monitoreo y argumentario de ventas",
        date     = "Mayo 2026  ·  v1.0  ·  Confidencial",
    )

    print("\nDocumentos generados en docs/")

