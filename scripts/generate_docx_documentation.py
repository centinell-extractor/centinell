from __future__ import annotations

from datetime import datetime
from pathlib import Path
import hashlib

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
OUT_PATH = ROOT / "docs" / "DOCUMENTACION_FUNCIONAL_TECNICA.docx"


def add_toc(document: Document) -> None:
    p = document.add_paragraph()
    run = p.add_run()

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    run._r.append(instr)

    fld_sep_run = p.add_run()
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_sep_run._r.append(fld_sep)
    fld_sep_run.add_text("Indice automatico. En Word: clic derecho -> Actualizar campo.")

    fld_end_run = p.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    fld_end_run._r.append(fld_end)


def file_sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def code_files() -> list[Path]:
    files: list[Path] = []

    for path in sorted((ROOT / "app").rglob("*.py")):
        if "__pycache__" not in path.parts:
            files.append(path)

    for path in sorted((ROOT / "app" / "static").glob("*")):
        if path.suffix in {".html", ".js", ".css", ".svg"}:
            files.append(path)

    for path in sorted((ROOT / "scripts").glob("*.sql")):
        files.append(path)

    for path in sorted((ROOT / "docs").glob("*.md")):
        files.append(path)

    for rel in [
        ".env.dev",
        ".env.example",
        ".gitignore",
        "samples/extract_body.json",
        "samples/extract_test_body.json",
    ]:
        p = ROOT / rel
        if p.exists():
            files.append(p)

    return files


def add_code_block(document: Document, content: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(content)
    run.font.name = "Consolas"
    run.font.size = document.styles["Normal"].font.size


def main() -> None:
    document = Document()

    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    document.add_heading("Centinell - Documentacion Funcional y Tecnica", level=0)
    document.add_paragraph(f"Fecha de generacion: {now}")
    document.add_paragraph(f"Workspace: {ROOT}")

    document.add_heading("Indice", level=1)
    add_toc(document)

    document.add_page_break()

    document.add_heading("1. Documento Funcional", level=1)
    document.add_paragraph(
        "Este apartado describe el comportamiento funcional del backend Centinell, "
        "los casos de uso principales y el estado validado en desarrollo."
    )

    document.add_heading("1.1 Objetivo", level=2)
    document.add_paragraph(
        "Proveer un flujo completo para extraccion de datos de facturas y documentos: "
        "configuracion de prompts, extraccion asistida por LLM, validacion humana, "
        "procesamiento en lote mediante colecciones y exportacion de resultados."
    )

    document.add_heading("1.2 Casos de uso cubiertos", level=2)
    document.add_paragraph("- Crear configuraciones de extraccion (POST /prompt-configs/).")
    document.add_paragraph("- Listar configuraciones existentes (GET /prompt-configs/).")
    document.add_paragraph("- Ejecutar extraccion real con config persistida (POST /extract/).")
    document.add_paragraph("- Guardar extracciones ligadas a una coleccion (collection_id opcional en POST /extract/).")
    document.add_paragraph("- Crear y listar colecciones para procesamiento batch (POST /collections/, GET /collections/).")
    document.add_paragraph("- Consultar extracciones por coleccion (GET /collections/{collection_id}/extractions).")
    document.add_paragraph("- Exportar colecciones a Excel (GET /collections/{collection_id}/export/xlsx).")
    document.add_paragraph("- Exportar extracciones por filtros en csv/xlsx/json (GET /extractions/export/bulk).")
    document.add_paragraph("- Exportar extraccion individual a Excel (GET /extractions/{id}/export/xlsx).")
    document.add_paragraph("- Probar extraccion con variables ad hoc (POST /extract-test/).")
    document.add_paragraph("- Ver prompt final de referencia (GET /test-template/).")

    document.add_heading("1.3 Estado validado", level=2)
    document.add_paragraph("- GET /prompt-configs/ -> 200 OK")
    document.add_paragraph("- POST /prompt-configs/ -> 201 Created")
    document.add_paragraph("- GET /collections/?limit=5 -> 200 OK")
    document.add_paragraph("- GET /extractions/?limit=5 -> 200 OK")
    document.add_paragraph("- GET /health -> 200 OK")
    document.add_paragraph(
        "- Conexion a Supabase estable mediante Session Pooler en entorno actual"
    )
    document.add_paragraph(
        "- Flujo UI validado: Nueva extraccion + Colecciones + Historial + exportaciones"
    )

    document.add_page_break()

    document.add_heading("2. Documento Tecnico", level=1)
    document.add_paragraph(
        "Este apartado incluye arquitectura, configuracion y el codigo fuente completo "
        "del estado actual del proyecto."
    )

    document.add_heading("2.1 Stack tecnico", level=2)
    document.add_paragraph("- FastAPI")
    document.add_paragraph("- SQLAlchemy async + asyncpg")
    document.add_paragraph("- Pydantic v2")
    document.add_paragraph("- httpx async")
    document.add_paragraph("- OpenAI Chat Completions")
    document.add_paragraph("- python-docx (generacion de esta documentacion)")
    document.add_paragraph("- openpyxl (exportacion XLSX)")
    document.add_paragraph("- Frontend vanilla: HTML + CSS + JavaScript")

    document.add_heading("2.2 Flujo tecnico resumido", level=2)
    document.add_paragraph("1) Carga de entorno en app/config.py con load_dotenv(override=True).")
    document.add_paragraph("2) Creacion de engine async y sesiones en app/db/connection.py.")
    document.add_paragraph("3) Routers FastAPI registrados en app/main.py.")
    document.add_paragraph("4) Extraccion via servicios de template, LLM y validador.")
    document.add_paragraph("5) Soporte de colecciones en BD (tabla collections + FK collection_id en extractions).")
    document.add_paragraph("6) Exportaciones: individual XLSX, bulk CSV/XLSX/JSON y XLSX por coleccion.")
    document.add_paragraph("7) Frontend con vista Colecciones para procesamiento batch desde multiples archivos.")

    document.add_heading("2.3 Codigo y configuracion completa", level=2)
    document.add_paragraph(
        "Se incluyen a continuacion todos los archivos de codigo y configuracion "
        "detectados en el proyecto al momento de generar este documento."
    )

    files = code_files()

    for path in files:
        rel = path.relative_to(ROOT)
        text = path.read_text(encoding="utf-8", errors="replace")

        document.add_heading(f"Archivo: {rel.as_posix()}", level=3)
        stat = path.stat()
        document.add_paragraph(
            f"Tamano: {stat.st_size} bytes | Ultima modificacion: "
            f"{datetime.fromtimestamp(stat.st_mtime).isoformat(timespec='seconds')}"
        )
        document.add_paragraph("Contenido:")
        add_code_block(document, text)

    document.add_page_break()

    document.add_heading("3. Referencias", level=1)
    document.add_paragraph(
        "Inventario de archivos incluidos (ruta relativa + hash SHA-256 para trazabilidad)."
    )

    for idx, path in enumerate(files, start=1):
        rel = path.relative_to(ROOT).as_posix()
        digest = file_sha256(path)
        document.add_paragraph(f"[{idx}] {rel} | sha256: {digest}")

    document.add_paragraph(
        "Nota: si el indice no aparece al abrir, actualizar campos en Word "
        "(clic derecho sobre el indice -> Actualizar campo)."
    )

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUT_PATH)
    print(str(OUT_PATH))


if __name__ == "__main__":
    main()
